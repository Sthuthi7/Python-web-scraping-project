import requests,json,bs4,os,time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import ElementClickInterceptedException,NoSuchElementException

from tkinter import *
from tkinter import ttk
import tkinter as tk
from tkinter import filedialog as fd
import nltk
nltk.download('punkt')
from newspaper import Article
from fpdf import FPDF
import textwrap
import docx,openpyxl,PyPDF2
from openpyxl.styles import Font

news_screen=Tk()#create a window
news_screen.title("NEWS")
news_screen.geometry("900x600")

canvas = Canvas(news_screen, width = 900, height = 600)  #add image    
canvas.pack(fill = "both", expand = True)      
img = PhotoImage(file="news.png")      
canvas.create_image(10,0,anchor="nw", image=img)

btn1 = Button(news_screen, text="FLASH NEWS",bg="orange", command=lambda:quick_field(),width=20,height=4)

btn1.place(x=170,y=250)

btn2 = Button(news_screen, text="DETAILED NEWS",bg="orange", command=lambda:detailed_field(),width=20,height=4)
btn2.place(x=500,y=250)

def Main_page(window):#to destroy the window once back button is pressed
       window.destroy()

def quick_field():#in this method we make use of selenium to extract news from website
    window = Toplevel(news_screen)
    window.title("NEWS")
    window.geometry("1200x650")
    canvas1 = Canvas(window, width = 1200, height = 650)      
    canvas1.pack(fill = "both", expand = True)    
    bg = PhotoImage(file="news3.png")      
    canvas1.create_image(140,0, anchor="nw", image=bg)  
    window.title("NEWS")
    lst=[]
     
    b=webdriver.Chrome()           #chrome webdriver
    b.maximize_window()
    b.get('https://www.ndtv.com/')
    try:
        try:
         pop1=b.find_element(By.CSS_SELECTOR,"#__cricketsubscribe > div.noti_wrap")
         pop=b.find_element(By.LINK_TEXT,"No Thanks")
         pop.click()
        except:
          print("No popup")
               
        news1=b.find_element(By.LINK_TEXT,"LATEST")
        news1.click()
        hl1='----LATEST----'     #get all the headlines from the different tabs in ndtv website
        lst.append(hl1)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems1=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(1) > div.news_Itm-cont > h2 > a')
        hl2=elems1[0].getText()
        lst.append(hl2)
        elems2=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(2) > div.news_Itm-cont > h2 > a')
        hl3=elems2[0].getText()
        lst.append(hl3)
        elems3=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(3) > div.news_Itm-cont > h2 > a')
        hl4=elems3[0].getText()
        lst.append(hl4)

        news2=b.find_element(By.LINK_TEXT,"INDIA")
        news2.click()
        hl5='----INDIA----'
        lst.append(hl5)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems4=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(1) > div.news_Itm-cont > h2 > a')
        hl6=elems4[0].getText()
        lst.append(hl6)
        elems5=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(2) > div.news_Itm-cont > h2 > a')
        hl7=elems5[0].getText()
        lst.append(hl7)
        elems6=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(3) > div.news_Itm-cont > h2 > a')
        hl8=elems6[0].getText()
        lst.append(hl8)

        news3=b.find_element(By.LINK_TEXT,"OPINION")
        news3.click()
        hl9='----OPINION----'
        lst.append(hl9)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems7=bs.select('body > div.content > div > div > section > div.row.s-lmr > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg.blog-opn > div > div:nth-child(1) > div.news_Itm-cont > h2 > a')
        hl10=elems7[0].getText()
        lst.append(hl10)
        elems8=bs.select('body > div.content > div > div > section > div.row.s-lmr > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg.blog-opn > div > div:nth-child(2) > div.news_Itm-cont > h2 > a')
        hl11=elems8[0].getText()
        lst.append(hl11)
        elems9=bs.select('body > div.content > div > div > section > div.row.s-lmr > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg.blog-opn > div > div:nth-child(3) > div.news_Itm-cont > h2 > a')
        hl12=elems9[0].getText()
        lst.append(hl12)

        news4=b.find_element(By.LINK_TEXT,"WORLD")
        news4.click()
        hl13='----WORLD----'
        lst.append(hl13)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems10=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(1) > div.news_Itm-cont > h2 > a')
        hl14=elems10[0].getText()
        lst.append(hl14)
        elems11=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(2) > div.news_Itm-cont > h2 > a')
        hl15=elems11[0].getText()
        lst.append(hl15)
        elems12=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(3) > div.news_Itm-cont > h2 > a')
        hl16=elems12[0].getText()
        lst.append(hl16)    

        news5=b.find_element(By.LINK_TEXT,"CRICKET")
        news5.click()
        hl17='----CRICKET----'
        lst.append(hl17)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems13=bs.select('body > div.bdy_ovr-wrp > div.section_two > div > div > div > div:nth-child(1) > div.vjl-sm-12.vjl-md-6.vjl-xl-5 > ul > li:nth-child(1) > div > div.crd_txt-wrp > h3 > a')
        hl18=elems13[0].getText()
        lst.append(hl18)
        elems14=bs.select('body > div.bdy_ovr-wrp > div.section_two > div > div > div > div:nth-child(1) > div.vjl-sm-12.vjl-md-6.vjl-xl-5 > ul > li:nth-child(2) > div > div.crd_txt-wrp > h3 > a')
        hl19=elems14[0].getText()
        lst.append(hl19)
        elems15=bs.select('body > div.bdy_ovr-wrp > div.section_two > div > div > div > div:nth-child(1) > div.vjl-sm-12.vjl-md-6.vjl-xl-5 > ul > li:nth-child(3) > div > div.crd_txt-wrp > h3 > a')
        hl20=elems15[0].getText()
        lst.append(hl20)

        news6=b.find_element(By.LINK_TEXT,"BUSINESS")
        news6.click()
        hl21='----BUSINESS----'
        lst.append(hl21)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems16=bs.select('body > div.newcont > div.newcont1 > div > div.newcont3 > div:nth-child(1) > div.lhs > div:nth-child(2) > div.l_lhs.topstories-cont > div.widcont_topstories > div.ll_rhs > div > ul > li:nth-child(1) > p > a')
        hl22=elems16[0].getText()
        lst.append(hl22)
        elems17=bs.select('body > div.newcont > div.newcont1 > div > div.newcont3 > div:nth-child(1) > div.lhs > div:nth-child(2) > div.l_lhs.topstories-cont > div.widcont_topstories > div.ll_rhs > div > ul > li:nth-child(2) > p > a')
        hl23=elems17[0].getText()
        lst.append(hl23)
        elems18=bs.select('body > div.newcont > div.newcont1 > div > div.newcont3 > div:nth-child(1) > div.lhs > div:nth-child(2) > div.l_lhs.topstories-cont > div.widcont_topstories > div.ll_rhs > div > ul > li:nth-child(3) > p > a')
        hl24=elems18[0].getText()
        lst.append(hl24)
        
        news7=b.find_element(By.LINK_TEXT,"MONEY")
        news7.click()
        hl25='----MONEY----'
        lst.append(hl25)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems19=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(1) > div.news_Itm-cont > h2 > a')
        hl26=elems19[0].getText()
        lst.append(hl26)
        elems20=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(2) > div.news_Itm-cont > h2 > a')
        hl27=elems20[0].getText()
        lst.append(hl27)
        elems21=bs.select('body > div.content > div > div > section > div.row.s-lmr.mt-10 > article > div > div > div > div.sp-cn.ins_storybody.lstng_Pg > div > div:nth-child(3) > div.news_Itm-cont > h2 > a')
        hl28=elems21[0].getText()
        lst.append(hl28)

        news8=b.find_element(By.LINK_TEXT,"TECH")
        news8.click()
        hl29='----TECH----'
        lst.append(hl29)
        url=b.current_url
        res=requests.get(url)
        bs=bs4.BeautifulSoup(res.text,features="lxml")
        elems22=bs.select('body > div.wrapper > div.row.gray_bg > div.container.padding_t25.clearfix > div.lhs_gray_section > div.recent_news_widget > ul > li:nth-child(2) > a')
        hl30=elems22[0].getText()
        lst.append(hl30)
        elems23=bs.select('body > div.wrapper > div.row.gray_bg > div.container.padding_t25.clearfix > div.lhs_gray_section > div.recent_news_widget > ul > li:nth-child(3) > a')
        hl31=elems23[0].getText()
        lst.append(hl31)
        elems24=bs.select('body > div.wrapper > div.row.gray_bg > div.container.padding_t25.clearfix > div.lhs_gray_section > div.recent_news_widget > ul > li:nth-child(4) > a')
        hl32=elems24[0].getText()
        lst.append(hl32)
       
    except IndexError:            #exception handling
       print('Index out of bound.')
    except ElementClickInterceptedException:          #incase there was some obstruction
        time.sleep(4)
    except NoSuchElementException:        #spelling error making this code not work as expected
        print("No Such Element Exception Handled")
    b.quit()
    file=open('adpnews.txt','w',encoding='utf-8')
    for i in lst:
           file.write(i+'\n\n\n')
    file.close()
    def toPDF():                 #convert to pdf
        # converting text to pdf file
        f=open("adpnews.txt","r",encoding='latin-1')
        text=f.read()
        f.close()
        pdf=FPDF(orientation='P', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_font("Arial",size=15)
        splitted = text.split('\n')
        for line in splitted:
          if line.startswith('----OPINION----'):
                 pdf.add_page()
          elif line.startswith('----CRICKET----'):
                 pdf.add_page()
          elif line.startswith('----MONEY----'):
                 pdf.add_page()
          lines = textwrap.wrap(line,75)
          if len(lines) == 0:
            pdf.ln()
          for wrap in lines:
            pdf.cell(0, 10,wrap, ln=1)
        pdf.output("quicknews.pdf")

        infile=open('quicknews.pdf','rb')
        pdfreader=PyPDF2.PdfFileReader(infile)
        pagenums=pdfreader.numPages
        wmfile=open('flash news.pdf',"rb")#merging watermark with pdf
        wmreader=PyPDF2.PdfFileReader(wmfile)
        pdfwriter=PyPDF2.PdfFileWriter()
        outfile=open("mergedQuickNewsPdf.pdf","wb")
        for page in range(int(pagenums)):
           p=pdfreader.getPage(page)
           wmpage=wmreader.getPage(0)
           p.mergePage(wmpage)
           pdfwriter.addPage(p)
           pdfwriter.write(outfile)
        outfile.close()
        infile.close()
        wmfile.close()
        os.startfile(r'C:\Users\LENOVO\AppData\Local\Programs\Python\Python39\mergedQuickNewsPdf.pdf')
    def toDocs(): #convert to docs
        doc =docx.Document()
        f=open("adpnews.txt","r",encoding='latin-1')
        text=f.read()
        f.close()
        lines=text.split('\n')
        for line in lines:
            if line.startswith('----LATEST----'):
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----INDIA----'):
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----OPINION----'):
              doc.add_page_break()
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----WORLD----'):
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----CRICKET----'):
              doc.add_page_break()
              para=doc.add_paragraph()
              para.add_run(line).bold=True  
            elif line.startswith('----BUSINESS----'):
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----MONEY----'):
              doc.add_page_break()
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            elif line.startswith('----TECH----'):
              para=doc.add_paragraph()
              para.add_run(line).bold=True
            else:
                   doc.add_paragraph(line)
        doc.save("quicknews.docx")
        os.startfile(r'C:\Users\LENOVO\AppData\Local\Programs\Python\Python39\quicknews.docx')
    
    lab=Label(window,text ="CONVERT TO....",width="30",height="2",font=("Calibri",20),bg="white")
    lab.place(x=340,y=100)
    btn1 = Button(window, text="PDF",bg="pink", command=toPDF,width=25,height=3)
    btn1.place(x=460,y=250)
    btn2= Button(window, text="DOCS",bg="pink", command=toDocs,width=25,height=3)
    btn2.place(x=460,y=380)
    def browseFiles():#to display the files created
     fd.askopenfilenames(filetypes=[("files",".pdf .docx")])
    explorebtn = Button(window,text = "Browse Files",bg="orange",command = browseFiles,width=30,height=4)
    explorebtn.place(x=450,y=500)
    Button(window, text="BACK",bg="red",fg="black",activebackground = "pink",width=15, height=2,command=lambda:Main_page(window)).place(x=1370,y=0)
    window.mainloop()
def detailed_field():#news extracted using web api called newsapi
    window = Toplevel(news_screen)
    canvas2 = Canvas(window, width = 1200, height = 650)      
    canvas2.pack(fill = "both", expand = True)    
    bg = PhotoImage(file="news2.png")      
    canvas2.create_image(500,70, anchor="nw", image=bg) 
    window.title("NEWS")
    window.geometry("1200x650")
    lab=Label(window,text ="ENTER THE NUMBER OF ARTICLES YOU WANT TO READ",width="50",height="2",font=("Calibri",15),bg="gray")
    lab.place(x=60,y=80)
    number=IntVar()
    num=Entry(window,textvariable=number,font=("Calibri",20),bg="pink")
    num.place(x = 250,y = 250,width=200,height=50)
    txt=Text(window,height=2,width=45,font=("Calibri",12))
    info="""(GIVE THE COUNT OF ARTICLES AND THEN SELECT THE   CONVERT OPTIONS)"""
    txt.insert(tk.END,info)
    txt.place(x=160,y=370)
    def click(btn1,btn2,btn3):
       btn1['state']="normal"
       btn2['state']="normal"
       btn3['state']="normal"
    
    def toPDF():#convert to pdf
     nonlocal num
     n=num.get()
     res=requests.get('https://newsapi.org/v2/top-headlines?country=in&apiKey=66658f95b8ef495998b2ab50de8f4281')
     info=json.loads(res.text)
     data=info['articles']
     lst=[]
     for i in range(int(n)):
        url =data[i]['url']
        article = Article(url)
        article.download()
        article.parse()
        article.nlp()
        lst.append('TITLE:\n'+data[i]['title'])
        if data[i]['description']!=None:
           lst.append('\n\nDESCRIPTION:\n\n'+data[i]['description'])
        if data[i]['urlToImage']!=None:
         lst.append('\n\n---IMAGE LINK---:\n\n'+data[i]['urlToImage'])
        lst.append('\n\n\nSUMMARY:\n'+article.summary)
        lst.append("\n\n\n\nARTICLE'S TEXT:\n"+article.text)
        lst.append("\n\n(ARTICLE"+str(i+1)+")")
        lst.append("\n\n\n\n")
     with open("random.txt","w",encoding='utf-8') as file:
         for l in lst:
           file.writelines(l)
     file.close()
     f=open("random.txt","r",encoding='latin-1')
     text=f.read()
     f.close()
     pdf=FPDF(orientation='P', unit='mm', format='A4')#writing from textfile to pdf
     pdf.set_font("Arial",size=15)
     splitted=text.split('\n')
     for line in splitted:
            if line.startswith('https://'):
              pdf.cell(0, 10,line, ln=1)
            else:
             if line.startswith('TITLE:') :
               pdf.add_page()
             lines=textwrap.wrap(line,75)
             if len(lines)==0:
               pdf.ln()
             for wrap in lines:
              pdf.cell(0, 10, wrap, ln=1)
     pdf.output("detailednews.pdf")

     infile=open('detailednews.pdf','rb')#used PyPDF2 module to merge two pdfs
     pdfreader=PyPDF2.PdfFileReader(infile)
     pagenums=pdfreader.numPages
     wmfile=open('news.pdf',"rb")
     wmreader=PyPDF2.PdfFileReader(wmfile)
     pdfwriter=PyPDF2.PdfFileWriter()
     outfile=open("mergednewspdf.pdf","wb")
     for page in range(int(pagenums)):
       p=pdfreader.getPage(page)
       wmpage=wmreader.getPage(0)
       p.mergePage(wmpage)
       pdfwriter.addPage(p)
       pdfwriter.write(outfile)
     outfile.close()
     infile.close()
     wmfile.close()
     os.startfile(r'C:\Users\LENOVO\AppData\Local\Programs\Python\Python39\mergednewspdf.pdf')
    def toDoc():#convert to docs
        n=num.get()
        res=requests.get('https://newsapi.org/v2/top-headlines?country=in&apiKey=66658f95b8ef495998b2ab50de8f4281')
        info=json.loads(res.text)
        data=info['articles']
        lst=[]
        for i in range(int(n)):
          url =data[i]['url']
          article = Article(url)
          article.download()
          article.parse()
          article.nlp()
          lst.append('TITLE:\n'+data[i]['title'])
          if data[i]['description']!=None:
             lst.append('\n\n\nDESCRIPTION:\n'+data[i]['description']+"...")
          lst.append('\n\n\n\nSUMMARY:\n'+article.summary)
          lst.append("\n\n\n\nARTICLE'S TEXT:\n"+article.text)
          lst.append("\n\n\n\n")
        with open("random.txt","w",encoding='utf-8') as file:
             for l in lst:
                file.writelines(l)
        file.close()
        doc=docx.Document()
        f=open("random.txt","r",encoding='latin-1')
        text=f.read()
        f.close()
        lines=text.split('\n')
        c=0
        for line in lines:
            if line.startswith('TITLE:'):
              if c>0:
                doc.add_page_break()
              c+=1
              para=doc.add_paragraph()
              b1=para.add_run(line)
              b1.bold=True
              b1.underline=True
            elif line.startswith('DESCRIPTION:'):
              para=doc.add_paragraph()
              b2=para.add_run(line)
              b2.bold=True
              b2.underline=True
            elif line.startswith('SUMMARY:'):
              para=doc.add_paragraph()
              b3=para.add_run(line)
              b3.bold=True
              b3.underline=True
            elif line.startswith("ARTICLE'S TEXT:"):
              para=doc.add_paragraph()
              b4=para.add_run(line)
              b4.bold=True
              b4.underline=True
            else:
                   doc.add_paragraph(line)
                   
        doc.save("detailednews.docx")
        os.startfile(r'C:\Users\LENOVO\AppData\Local\Programs\Python\Python39\detailednews.docx')
    def toEXCEL():#convert to excel
          nonlocal num
          n=num.get()
          res=requests.get('https://newsapi.org/v2/top-headlines?country=in&apiKey=66658f95b8ef495998b2ab50de8f4281')
          info=json.loads(res.text)
          data=info['articles']
          lst=[]
          mainlst=[]
          for i in range(int(n)):
             lst.append(data[i]['source']['name'])
             pub=data[i]['publishedAt'].split('T')
             lst.append(pub[0])
             st=pub[1].split('Z')
             lst.append(st[0])
             mainlst+=[lst]
             lst=[]
          wb=openpyxl.Workbook()
          sheet=wb["Sheet"]
          font1=Font(size=10,bold=True)
          sheet.append(['NAME','PUBLISHED DATE','PUBLISHED TIME'])
          sheet['A1'].font=font1
          sheet['B1'].font=font1
          sheet['C1'].font=font1
          sheet.column_dimensions['A'].width=20
          sheet.column_dimensions['B'].width=20
          sheet.column_dimensions['C'].width=20
          for row in mainlst:
                 sheet.append(row)#append the contents of list into sheet
          wb.save("detailednews.xlsx")
          os.startfile(r'C:\Users\LENOVO\AppData\Local\Programs\Python\Python39\detailednews.xlsx')
    
    lab=Label(window,text ="CONVERT TO....",width="30",height="2",font=("Calibri",15),bg="gray")
    lab.place(x=750,y=90)
    btn1 = Button(window, text="PDF",bg="pink", command=toPDF,width=25,height=3,state="disabled")
    btn1.place(x=850,y=180)
    btn2 = Button(window, text="DOCS",bg="pink", command=toDoc,width=25,height=3,state="disabled")
    btn2.place(x=850,y=270)
    btn3 = Button(window, text="EXCEL",bg="pink", command=toEXCEL,width=25,height=3,state="disabled")
    btn3.place(x=850,y=360)
    txt3=Text(window,height=3,width=45,font=("Calibri",13))
    info3="""EXCEL will store the name,Published date and published time of the  website of the articles"""
    txt3.insert(tk.END,info3)
    txt3.place(x=720,y=430)
    def browseFiles():
       fd.askopenfilenames(filetypes=[("files",".xlsx .pdf .docx")])
    subbtn = Button(window, text="SUBMIT",bg="orange", command=lambda:click(btn1,btn2,btn3),width=43,height=3)
    subbtn.place(x=200,y=430)
    explorebtn = Button(window,text = "Browse Files",bg="orange",command = browseFiles,width=30,height=4)
    explorebtn.place(x=830,y=510)
    Button(window, text="BACK",bg="red",fg="black",activebackground = "pink",width=15, height=2,command=lambda:Main_page(window)).place(x=1370,y=0)
    window.mainloop()#button to return to main window
